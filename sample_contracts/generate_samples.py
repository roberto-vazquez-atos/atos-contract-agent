"""Generate sample Word contracts for demo purposes."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def create_sample_nda():
    doc = Document()
    doc.add_heading("NON-DISCLOSURE AGREEMENT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Non-Disclosure Agreement (\"Agreement\") is entered into as of January 15, 2025, "
        "between Atos SE, a corporation organized under the laws of France, with its principal "
        "place of business at 80 Quai Voltaire, 95870 Bezons, France (\"Disclosing Party\"), "
        "and TechVenture GmbH, a corporation organized under the laws of Germany, located at "
        "Hauptstraße 12, 10115 Berlin, Germany (\"Receiving Party\")."
    )

    doc.add_heading("1. Definition of Confidential Information", level=1)
    doc.add_paragraph(
        "\"Confidential Information\" means any and all technical and non-technical information "
        "disclosed by Disclosing Party to Receiving Party, including but not limited to: "
        "trade secrets, inventions, ideas, processes, formulas, source code, data, programs, "
        "business plans, customer lists, financial information, and any other proprietary information."
    )
    doc.add_paragraph(
        "Confidential Information does NOT include information that: (a) is or becomes publicly "
        "available through no fault of Receiving Party; (b) was rightfully known to Receiving "
        "Party prior to disclosure; (c) is independently developed by Receiving Party without "
        "use of Confidential Information."
    )

    doc.add_heading("2. Obligations of Receiving Party", level=1)
    doc.add_paragraph(
        "Receiving Party agrees to: (a) hold Confidential Information in strict confidence; "
        "(b) not disclose Confidential Information to any third party without prior written "
        "consent; (c) use Confidential Information solely for evaluating a potential business "
        "relationship between the parties; (d) limit access to those employees with a need to know."
    )

    doc.add_heading("3. Term", level=1)
    doc.add_paragraph(
        "This Agreement shall commence on the Effective Date and continue for a period of two (2) "
        "years. Obligations of confidentiality shall survive termination of this Agreement for "
        "a period of one (1) year."
    )

    doc.add_heading("4. Return of Information", level=1)
    doc.add_paragraph(
        "Upon written request by Disclosing Party, Receiving Party shall promptly return or "
        "destroy all Confidential Information and certify such destruction in writing."
    )

    doc.add_heading("5. No License", level=1)
    doc.add_paragraph(
        "Nothing in this Agreement grants Receiving Party any rights in or to Confidential "
        "Information except as expressly set forth herein."
    )

    doc.add_heading("6. Governing Law", level=1)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with the laws of France, "
        "without regard to its conflict of law provisions."
    )

    doc.add_heading("7. Signatures", level=1)
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.")
    doc.add_paragraph("\nAtos SE\nBy: ___________________________\nName:\nTitle:\nDate:")
    doc.add_paragraph("\nTechVenture GmbH\nBy: ___________________________\nName:\nTitle:\nDate:")

    path = os.path.join(os.path.dirname(__file__), "sample_nda.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_sample_sow():
    doc = Document()
    doc.add_heading("STATEMENT OF WORK", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Cloud Migration Services — Project Aurora", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Statement of Work (\"SOW\") is entered into as of March 1, 2025, between "
        "Atos IT Solutions and Services, Inc. (\"Service Provider\") and RetailCorp SA "
        "(\"Client\"), and is incorporated into the Master Services Agreement dated February 1, 2025."
    )

    doc.add_heading("1. Project Scope", level=1)
    doc.add_paragraph(
        "Service Provider shall deliver cloud migration services for Client's legacy ERP system "
        "to Microsoft Azure. The scope includes: infrastructure assessment, migration planning, "
        "data migration, testing, and go-live support. Out of scope: end-user training, "
        "post-migration support beyond 30 days, and third-party license procurement."
    )

    doc.add_heading("2. Deliverables", level=1)
    deliverables = [
        "D1 — Infrastructure Assessment Report (due: April 15, 2025)",
        "D2 — Migration Plan and Architecture Design (due: May 1, 2025)",
        "D3 — Migrated Environment (dev/staging) (due: June 30, 2025)",
        "D4 — Production Migration and Go-Live (due: August 15, 2025)",
        "D5 — Post-Migration Validation Report (due: September 15, 2025)",
    ]
    for d in deliverables:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("3. Payment Terms", level=1)
    doc.add_paragraph(
        "Total contract value: EUR 450,000. Payment schedule: 20% upon SOW signature, "
        "20% upon delivery of D2, 30% upon delivery of D3, 30% upon delivery of D5. "
        "Invoices payable within 45 days of receipt."
    )

    doc.add_heading("4. Roles and Responsibilities", level=1)
    doc.add_paragraph(
        "Service Provider: Project Manager (J. Martin), Lead Cloud Architect (S. Kumar), "
        "2x Migration Engineers. Client: IT Project Sponsor, IT Infrastructure Lead, "
        "ERP System Owner."
    )

    doc.add_heading("5. Governing Law", level=1)
    doc.add_paragraph(
        "This SOW shall be governed by the laws of France. Disputes shall be resolved by "
        "the Commercial Court of Paris."
    )

    path = os.path.join(os.path.dirname(__file__), "sample_sow.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_mclinic_sow_1():
    """SOW 1: M Clinic Predictive Readmission Model — well-structured, follows Atos standards."""
    doc = Document()
    doc.add_heading("STATEMENT OF WORK", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Predictive Patient Readmission Model — Project PRISM", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Statement of Work (\"SOW\") is entered into as of February 3, 2025, between "
        "Atos IT Solutions and Services, S.A.U. (\"Service Provider\") and M Clinic, S.A., "
        "a healthcare institution organized under the laws of Spain, with its principal place "
        "of business at Carrer del Rosselló 76, 08029 Barcelona, Spain (\"Client\"), and is "
        "incorporated into the Master Services Agreement (\"MSA\") dated January 10, 2025 "
        "(MSA Reference: MSA-MCLINIC-2025-001)."
    )

    doc.add_heading("1. Project Scope", level=1)
    doc.add_paragraph(
        "Service Provider shall design, develop, validate, and deploy a machine learning model "
        "to predict 30-day hospital readmission risk for patients discharged from Client's "
        "cardiology and internal medicine units. The model will be trained on anonymised "
        "electronic health record (EHR) data provided by Client and integrated into Client's "
        "existing clinical decision support platform (Cerner Millennium)."
    )
    doc.add_paragraph("The following are explicitly OUT OF SCOPE:")
    out_of_scope = [
        "Procurement or licensing of third-party data sources",
        "Modifications to the Cerner Millennium platform beyond agreed API integration points",
        "Clinical validation studies or regulatory submissions (CE marking, FDA clearance)",
        "End-user training beyond two (2) train-the-trainer sessions",
        "Ongoing model maintenance after the warranty period defined in Section 8",
    ]
    for item in out_of_scope:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Deliverables and Acceptance Criteria", level=1)
    doc.add_paragraph(
        "Each deliverable below is subject to formal written acceptance by Client within "
        "ten (10) business days of delivery. Acceptance shall not be unreasonably withheld. "
        "If Client does not respond within ten (10) business days, the deliverable shall be "
        "deemed accepted."
    )
    deliverables = [
        ("D1", "Data Assessment and Feasibility Report", "March 14, 2025",
         "Report documents data quality findings, feature candidates, and a confirmed feasibility recommendation."),
        ("D2", "Model Architecture and Development Plan", "April 4, 2025",
         "Document specifies model approach, validation strategy, KPIs, and updated project timeline."),
        ("D3", "Baseline Model (v1.0) with Performance Report", "June 6, 2025",
         "Model achieves minimum AUC-ROC ≥ 0.78 on held-out test set. Performance report provided."),
        ("D4", "Integrated Model in Staging Environment", "July 18, 2025",
         "End-to-end integration tested in Client's staging environment with sign-off from Client IT Lead."),
        ("D5", "Production Deployment and Go-Live", "September 5, 2025",
         "Model live in production. Monitoring dashboard active. Go-live report signed by both parties."),
        ("D6", "Post-Deployment Validation Report (90-day)", "December 5, 2025",
         "Report demonstrates sustained AUC-ROC ≥ 0.75 over 90 days of production operation."),
    ]
    for code, name, due, criteria in deliverables:
        doc.add_paragraph(f"{code} — {name} (Due: {due})", style="List Bullet")
        p = doc.add_paragraph(f"Acceptance criteria: {criteria}")
        p.paragraph_format.left_indent = Pt(36)

    doc.add_heading("3. Payment Terms", level=1)
    doc.add_paragraph(
        "Total contract value: EUR 320,000 (excluding VAT). Payment schedule linked to deliverable acceptance:"
    )
    payments = [
        "15% (EUR 48,000) — upon SOW execution",
        "15% (EUR 48,000) — upon acceptance of D2",
        "25% (EUR 80,000) — upon acceptance of D3",
        "25% (EUR 80,000) — upon acceptance of D5",
        "20% (EUR 64,000) — upon acceptance of D6",
    ]
    for p in payments:
        doc.add_paragraph(p, style="List Bullet")
    doc.add_paragraph(
        "All invoices are payable within thirty (30) calendar days of receipt. Late payments "
        "shall accrue interest at the European Central Bank base rate plus 8%, as per Directive 2011/7/EU."
    )

    doc.add_heading("4. Roles and Responsibilities", level=1)
    doc.add_paragraph("Service Provider shall provide:")
    sp_roles = [
        "Project Manager: single point of accountability for delivery and escalation",
        "Lead Data Scientist: responsible for model architecture and validation",
        "2× Senior ML Engineers: feature engineering, training pipeline, and API development",
        "Integration Engineer: Cerner Millennium API integration",
    ]
    for r in sp_roles:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_paragraph("Client shall provide:")
    client_roles = [
        "Clinical Project Sponsor: executive decision-making authority",
        "Data Access Manager: responsible for EHR data extraction and anonymisation",
        "IT Integration Lead: Cerner Millennium access and environment provisioning",
        "Clinical Expert (Cardiology/Internal Medicine): domain validation of model outputs",
    ]
    for r in client_roles:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("5. Change Management", level=1)
    doc.add_paragraph(
        "Any change to scope, timeline, or budget must be submitted via a written Change Request "
        "(\"CR\") using the template in Annex A. Service Provider shall assess and respond to each "
        "CR within five (5) business days with an impact assessment covering effort, cost, and "
        "schedule. No change shall be implemented without a CR approved in writing by both parties' "
        "designated project managers."
    )

    doc.add_heading("6. Intellectual Property", level=1)
    doc.add_paragraph(
        "All intellectual property developed exclusively for this engagement and funded entirely "
        "by Client under this SOW (\"Custom IP\") shall be owned by Client upon full payment. "
        "Service Provider retains ownership of all pre-existing intellectual property, tools, "
        "frameworks, and methodologies (\"Background IP\") used in delivery. Service Provider "
        "grants Client a perpetual, non-exclusive, royalty-free licence to use Background IP "
        "solely to the extent embedded in the Custom IP deliverables."
    )

    doc.add_heading("7. Data Protection and Confidentiality", level=1)
    doc.add_paragraph(
        "All personal data processed under this SOW shall be handled in accordance with "
        "Regulation (EU) 2016/679 (GDPR). Service Provider acts as Data Processor; Client acts "
        "as Data Controller. A Data Processing Agreement (\"DPA\") is attached as Annex B and "
        "forms part of this SOW. Both parties shall maintain the confidentiality of the other "
        "party's Confidential Information for the duration of this SOW and for five (5) years thereafter."
    )

    doc.add_heading("8. Warranty", level=1)
    doc.add_paragraph(
        "Service Provider warrants that deliverables will conform to agreed acceptance criteria "
        "for a period of ninety (90) days following acceptance of D5 (\"Warranty Period\"). "
        "During the Warranty Period, Service Provider shall remedy any non-conformance at no "
        "additional cost to Client within ten (10) business days of written notification."
    )

    doc.add_heading("9. Limitation of Liability", level=1)
    doc.add_paragraph(
        "Each party's aggregate liability under this SOW shall not exceed the total fees paid "
        "or payable under this SOW in the twelve (12) months preceding the event giving rise to "
        "the claim. Neither party shall be liable for indirect, consequential, or punitive damages."
    )

    doc.add_heading("10. Termination", level=1)
    doc.add_paragraph(
        "Either party may terminate this SOW for convenience upon thirty (30) calendar days' "
        "written notice. In the event of termination for convenience by Client, Client shall pay "
        "for all work completed and accepted up to the termination date, plus reasonable "
        "wind-down costs not to exceed EUR 15,000. Either party may terminate immediately upon "
        "written notice if the other party is in material breach and fails to cure within fifteen "
        "(15) business days of written notice."
    )

    doc.add_heading("11. Force Majeure", level=1)
    doc.add_paragraph(
        "Neither party shall be liable for failure or delay in performance to the extent caused "
        "by circumstances beyond its reasonable control, including natural disasters, acts of "
        "government, pandemics, or cyberattacks. The affected party shall notify the other within "
        "five (5) business days and shall use reasonable efforts to mitigate the impact."
    )

    doc.add_heading("12. Governing Law and Dispute Resolution", level=1)
    doc.add_paragraph(
        "This SOW shall be governed by the laws of France. Any dispute shall first be referred "
        "to the parties' respective senior management for resolution within thirty (30) days. "
        "Unresolved disputes shall be submitted to binding arbitration under the ICC Rules, "
        "with the seat of arbitration in Paris, France, conducted in English."
    )

    doc.add_heading("13. Signatures", level=1)
    doc.add_paragraph("IN WITNESS WHEREOF, the parties have executed this Statement of Work as of the date first written above.")
    doc.add_paragraph("\nAtos IT Solutions and Services, S.A.U.\nBy: ___________________________\nName:\nTitle:\nDate:")
    doc.add_paragraph("\nM Clinic, S.A.\nBy: ___________________________\nName:\nTitle:\nDate:")

    path = os.path.join(os.path.dirname(__file__), "mclinic_sow_readmission.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_mclinic_sow_2():
    """SOW 2: M Clinic Medical Imaging Classifier — inconsistent standards, missing clauses."""
    doc = Document()
    doc.add_heading("Statement of Work", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("AI-Based Radiology Image Classification System", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This document constitutes a Statement of Work agreed between Atos (hereinafter \"Atos\") "
        "and M Clinic (hereinafter sometimes referred to as \"the Hospital\" or \"the institution\"). "
        "Work shall commence on or around 1st April 2025 and is related to an AI imaging project. "
        "This SOW should be read alongside any previously signed agreements between the two organisations."
    )
    # NOTE: No MSA reference, inconsistent party names, vague start date

    doc.add_heading("Scope of Work", level=1)
    # NOTE: Not numbered, different heading style than SOW 1
    doc.add_paragraph(
        "Atos will build a deep learning classifier to identify and flag anomalies in chest X-ray "
        "and CT scan images uploaded to M Clinic's radiology PACS system (GE Centricity). The "
        "solution will cover the following pathologies: pneumonia, pleural effusion, pulmonary "
        "nodules (>6mm), and cardiomegaly. The Hospital is responsible for providing labelled "
        "training data and access to the PACS environment."
    )
    doc.add_paragraph(
        "Things NOT included: we won't be doing any work on MRI images in this phase. "
        "Integration with any systems other than GE Centricity is also out of scope."
    )
    # NOTE: Informal language, incomplete out-of-scope

    doc.add_heading("Milestones and Outputs", level=1)
    # NOTE: Uses "Milestones" instead of "Deliverables", no formal acceptance criteria
    milestones = [
        "Milestone 1 — Kick-off and data review: April 2025",
        "Milestone 2 — First model prototype: June 2025",
        "Milestone 3 — Refined model after clinical feedback: August 2025",
        "Milestone 4 — System deployed to production: October 2025",
    ]
    for m in milestones:
        doc.add_paragraph(m, style="List Bullet")
    doc.add_paragraph(
        "Outputs will be reviewed and signed off by M Clinic's radiology team. Specific "
        "performance targets to be agreed during the project."
    )
    # NOTE: No concrete acceptance criteria, performance targets undefined

    doc.add_heading("Fees and Payment", level=1)
    # NOTE: Different section name ("Fees and Payment" vs "Payment Terms")
    doc.add_paragraph(
        "The total cost for this engagement is EUR 290,000 plus applicable taxes. "
        "Invoices will be issued as follows: 20% at start, 30% at Milestone 2, 50% at Milestone 4. "
        "Payment is expected within 45 days of invoice date."
    )
    # NOTE: Net 45 vs Net 30 in SOW 1, vague milestone-to-payment linkage

    doc.add_heading("Team", level=1)
    doc.add_paragraph(
        "Atos will assign a project lead and the necessary technical staff to deliver the above. "
        "M Clinic will need to make available radiology staff for feedback sessions and an IT "
        "contact for system access."
    )
    # NOTE: No named roles, no commitment to specific headcount

    doc.add_heading("Intellectual Property", level=1)
    doc.add_paragraph(
        "The model and associated code produced during this engagement will belong to Atos. "
        "M Clinic will receive a licence to use the solution for internal clinical purposes. "
        "M Clinic may not resell or sublicence the solution to third parties."
    )
    # NOTE: IP ownership is REVERSED vs SOW 1 — Atos retains ownership here
    # This is a major inconsistency that the compare tool should flag

    doc.add_heading("Data and Privacy", level=1)
    doc.add_paragraph(
        "All patient data used for training must be fully anonymised by the Hospital before "
        "being shared with Atos. Atos will handle this data securely and delete it upon project "
        "completion. Both parties agree to maintain confidentiality of project information "
        "during the project."
    )
    # NOTE: No GDPR DPA mentioned, confidentiality duration not specified (vs 5 years in SOW 1)

    doc.add_heading("Termination", level=1)
    doc.add_paragraph(
        "Either party can end this agreement by giving 15 days notice in writing. "
        "In this case, M Clinic will pay for work completed up to that point."
    )
    # NOTE: 15 days vs 30 days in SOW 1, no wind-down cost cap, informal language

    doc.add_heading("Applicable Law", level=1)
    # NOTE: Different section name ("Applicable Law" vs "Governing Law and Dispute Resolution")
    doc.add_paragraph(
        "This SOW is governed by the laws of Spain. Any disputes will be handled by the "
        "competent courts of Barcelona."
    )
    # NOTE: Spain vs France — inconsistency with Atos standard and SOW 1

    doc.add_heading("Signatures", level=1)
    doc.add_paragraph("Agreed and signed on behalf of the parties:")
    doc.add_paragraph("\nAtos\nSigned: ___________________________\nName:\nDate:")
    doc.add_paragraph("\nM Clinic\nSigned: ___________________________\nName:\nDate:")
    # NOTE: Missing titles, informal "Signed" vs "By:", no full legal entity names

    path = os.path.join(os.path.dirname(__file__), "mclinic_sow_imaging.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_nda_fintech():
    """Reference NDA: Atos SE with FinanceApp Ltd (UK fintech) — all required clauses present."""
    doc = Document()
    doc.add_heading("NON-DISCLOSURE AGREEMENT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Non-Disclosure Agreement (\"Agreement\") is entered into as of 1 April 2025, "
        "between Atos SE, a société européenne organised under the laws of France, with its "
        "registered office at 80 Quai Voltaire, 95870 Bezons, France (\"Disclosing Party\"), "
        "and FinanceApp Ltd, a private limited company incorporated in England and Wales "
        "(Company No. 09871234), with its registered office at 1 Canada Square, London E14 5AB, "
        "United Kingdom (\"Receiving Party\"). "
        "Disclosing Party and Receiving Party are each referred to herein as a \"Party\" and "
        "collectively as the \"Parties\"."
    )

    doc.add_paragraph(
        "WHEREAS, the Parties wish to explore a potential strategic partnership relating to "
        "open-banking API integration services (the \"Purpose\"), and, in connection therewith, "
        "Disclosing Party may disclose Confidential Information to Receiving Party."
    )
    doc.add_paragraph(
        "NOW, THEREFORE, in consideration of the mutual covenants herein and for other good and "
        "valuable consideration, the receipt and sufficiency of which are hereby acknowledged, "
        "the Parties agree as follows."
    )

    doc.add_heading("1. Definition of Confidential Information", level=1)
    doc.add_paragraph(
        "\"Confidential Information\" means any and all information or data that has or could "
        "have commercial value or other utility in the business in which Disclosing Party is "
        "engaged, disclosed by Disclosing Party to Receiving Party, whether disclosed orally, "
        "in writing, electronically, or by any other means, including but not limited to: "
        "business plans, financial projections, client and prospect lists, pricing structures, "
        "technical architectures, software source code, product roadmaps, proprietary "
        "methodologies, and any information marked as \"Confidential\" or \"Proprietary\"."
    )
    doc.add_paragraph(
        "Confidential Information shall NOT include information that: (a) is or becomes part "
        "of the public domain through no act or omission of Receiving Party; (b) was in "
        "Receiving Party's lawful possession prior to disclosure without restriction on "
        "disclosure; (c) is independently developed by Receiving Party without use of or "
        "reference to Confidential Information, as evidenced by contemporaneous written records; "
        "or (d) is required to be disclosed by applicable law, regulation, or court order, "
        "provided that Receiving Party gives Disclosing Party prompt written notice and "
        "cooperates with Disclosing Party to seek a protective order."
    )

    doc.add_heading("2. Obligations of Receiving Party", level=1)
    doc.add_paragraph(
        "Receiving Party agrees to: (a) hold all Confidential Information in strict confidence "
        "and protect it with at least the same degree of care it uses to protect its own "
        "confidential information, but no less than reasonable care; "
        "(b) not disclose, copy, reproduce, or distribute Confidential Information to any third "
        "party without Disclosing Party's prior written consent; "
        "(c) use Confidential Information solely for the Purpose; "
        "(d) limit access to those of its employees, officers, and professional advisors who "
        "have a strict need to know for the Purpose, and who are bound by confidentiality "
        "obligations at least as protective as those set out herein; "
        "(e) immediately notify Disclosing Party upon becoming aware of any unauthorised "
        "disclosure or use of Confidential Information."
    )

    doc.add_heading("3. Term and Duration", level=1)
    doc.add_paragraph(
        "This Agreement shall commence on the date first written above and shall remain in "
        "full force for a period of three (3) years (the \"Term\"). Either Party may terminate "
        "this Agreement at any time upon thirty (30) days' written notice. "
        "The obligations of confidentiality shall survive expiry or termination of this "
        "Agreement for a further period of two (2) years."
    )

    doc.add_heading("4. Return or Destruction of Information", level=1)
    doc.add_paragraph(
        "Upon expiry or termination of this Agreement, or upon written request by Disclosing "
        "Party at any time, Receiving Party shall promptly — and in any case within ten (10) "
        "business days — return to Disclosing Party or securely destroy all tangible materials "
        "containing Confidential Information, and delete all electronic copies from its systems. "
        "Receiving Party shall certify such return or destruction in writing within five (5) "
        "business days of completion. A single archival copy may be retained solely for "
        "compliance purposes, subject to the ongoing confidentiality obligations herein."
    )

    doc.add_heading("5. No Licence Granted", level=1)
    doc.add_paragraph(
        "Nothing in this Agreement shall be construed as granting Receiving Party any right, "
        "licence, or interest in or to the Confidential Information, or any intellectual "
        "property rights of Disclosing Party, except the limited right to use Confidential "
        "Information solely for the Purpose as expressly set out herein."
    )

    doc.add_heading("6. Standard Exceptions", level=1)
    doc.add_paragraph(
        "Receiving Party shall not be in breach of this Agreement if disclosure of Confidential "
        "Information is required: (a) by a court of competent jurisdiction; (b) by a regulatory "
        "or governmental authority having jurisdiction over Receiving Party; or (c) by applicable "
        "law. In each such case, Receiving Party shall, to the extent permitted by law, give "
        "Disclosing Party prior written notice and shall disclose only that portion of the "
        "Confidential Information that is legally required to be disclosed."
    )

    doc.add_heading("7. Dispute Resolution", level=1)
    doc.add_paragraph(
        "In the event of any dispute arising out of or in connection with this Agreement, the "
        "Parties shall first attempt to resolve the matter amicably through good-faith "
        "negotiations between senior representatives of each Party within thirty (30) days of "
        "written notice of the dispute. If the matter is not resolved within that period, "
        "either Party may refer the dispute to the courts under Clause 8."
    )

    doc.add_heading("8. Governing Law and Jurisdiction", level=1)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with the laws of "
        "England and Wales, without regard to its conflict of laws principles. Each Party "
        "irrevocably submits to the exclusive jurisdiction of the courts of England and Wales "
        "for the resolution of any dispute arising out of or in connection with this Agreement."
    )

    doc.add_heading("9. General Provisions", level=1)
    doc.add_paragraph(
        "This Agreement constitutes the entire agreement between the Parties with respect to "
        "its subject matter and supersedes all prior oral and written agreements. "
        "This Agreement may only be amended by a written instrument signed by authorised "
        "representatives of both Parties. If any provision of this Agreement is held to be "
        "invalid or unenforceable, the remaining provisions shall continue in full force."
    )

    doc.add_heading("10. Signatures", level=1)
    doc.add_paragraph(
        "IN WITNESS WHEREOF, the Parties have executed this Non-Disclosure Agreement as of "
        "the date first written above."
    )
    doc.add_paragraph("\nAtos SE\nBy: ___________________________\nName:\nTitle:\nDate:")
    doc.add_paragraph("\nFinanceApp Ltd\nBy: ___________________________\nName:\nTitle:\nDate:")

    path = os.path.join(os.path.dirname(__file__), "nda_financeapp.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_nda_automotive():
    """Reference NDA: Atos SE with AutoDrive AG (automotive) — all required clauses present."""
    doc = Document()
    doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Mutual Non-Disclosure Agreement (\"Agreement\") is entered into as of 15 May 2025, "
        "between Atos SE, a société européenne organised under the laws of France, with its "
        "registered office at 80 Quai Voltaire, 95870 Bezons, France (\"Party A\"), "
        "and AutoDrive AG, a company organised under the laws of Germany, with its registered "
        "office at Maximilianstraße 4, 80539 Munich, Germany (\"Party B\"). "
        "This Agreement is entered into in connection with the Parties' joint evaluation of a "
        "potential collaboration on connected vehicle data analytics (the \"Purpose\"). "
        "Each Party may act as both Disclosing Party and Receiving Party under this Agreement."
    )

    doc.add_heading("1. Definition of Confidential Information", level=1)
    doc.add_paragraph(
        "\"Confidential Information\" means all information, in any form or medium, that a "
        "Party (the \"Disclosing Party\") discloses or makes available to the other Party "
        "(the \"Receiving Party\") in connection with the Purpose, including without limitation: "
        "vehicle telemetry architectures, sensor data schemas, proprietary algorithms, software "
        "components, business strategies, customer data models, financial projections, and "
        "personnel information. Information is Confidential Information whether or not it is "
        "marked or designated as such at the time of disclosure, if a reasonable person would "
        "understand it to be confidential in the context of the disclosure."
    )
    doc.add_paragraph(
        "Confidential Information excludes information that: (a) is or becomes generally "
        "available to the public other than as a result of a breach of this Agreement; "
        "(b) was already known to Receiving Party prior to disclosure as evidenced by written "
        "records predating the disclosure; (c) becomes known to Receiving Party from a third "
        "party who is not known to be under any confidentiality obligation with respect to "
        "such information; or (d) is independently developed by Receiving Party's personnel "
        "who had no access to the Confidential Information."
    )

    doc.add_heading("2. Obligations of Receiving Party", level=1)
    doc.add_paragraph(
        "Each Receiving Party shall: (a) use the Disclosing Party's Confidential Information "
        "solely for the Purpose; (b) protect Confidential Information with the same standard "
        "of care applied to its own most sensitive information, but in no event less than "
        "reasonable care; (c) not disclose Confidential Information to any person or entity "
        "other than its own employees, contractors, and professional advisors who need access "
        "for the Purpose and who are bound by confidentiality obligations no less stringent "
        "than those in this Agreement; (d) promptly notify the Disclosing Party of any actual "
        "or suspected unauthorised access, use, or disclosure of Confidential Information; and "
        "(e) take all reasonable steps to prevent unauthorised use or disclosure."
    )

    doc.add_heading("3. Term and Duration", level=1)
    doc.add_paragraph(
        "This Agreement takes effect on the date first written above and continues for a period "
        "of two (2) years (\"Initial Term\"). Either Party may extend this Agreement for an "
        "additional one (1) year by written notice given at least thirty (30) days before "
        "expiry. The confidentiality obligations in this Agreement shall survive expiry or "
        "termination and shall remain binding for a further three (3) years."
    )

    doc.add_heading("4. Return or Destruction of Information", level=1)
    doc.add_paragraph(
        "Within fifteen (15) business days of the expiry or termination of this Agreement, or "
        "upon written request by the Disclosing Party, the Receiving Party shall, at the "
        "Disclosing Party's election, either return all Confidential Information and all copies "
        "thereof, or securely destroy all such information and provide written certification of "
        "destruction. The foregoing obligation extends to electronic data stored on servers, "
        "cloud platforms, and portable devices."
    )

    doc.add_heading("5. No Licence Granted", level=1)
    doc.add_paragraph(
        "No licence or other right is granted to either Party under this Agreement, whether "
        "expressly, by implication, estoppel, or otherwise, in respect of any Confidential "
        "Information, intellectual property, patent, copyright, trademark, or trade secret. "
        "The disclosure of Confidential Information does not obligate either Party to enter "
        "into any further agreement or business relationship."
    )

    doc.add_heading("6. Standard Exceptions", level=1)
    doc.add_paragraph(
        "Each Party acknowledges that its obligations under this Agreement shall not apply "
        "to the extent that Confidential Information must be disclosed: (a) pursuant to a "
        "legally binding order of a court, arbitral tribunal, or regulatory authority; or "
        "(b) as required by any applicable law or regulation. In each such case, the "
        "Receiving Party shall give the Disclosing Party prompt advance written notice (to the "
        "extent legally permitted) and shall cooperate with the Disclosing Party in seeking "
        "appropriate confidential treatment for the disclosure."
    )

    doc.add_heading("7. Dispute Resolution", level=1)
    doc.add_paragraph(
        "Any dispute, controversy, or claim arising out of or relating to this Agreement, "
        "or its breach, termination, or validity, shall first be referred to the senior "
        "management of both Parties for good-faith resolution within twenty-one (21) calendar "
        "days. If not resolved within that period, the dispute shall be finally settled by "
        "arbitration under the Rules of Arbitration of the International Chamber of Commerce "
        "(ICC), with a sole arbitrator, the seat of arbitration in Frankfurt, Germany, and "
        "the language of proceedings English."
    )

    doc.add_heading("8. Governing Law", level=1)
    doc.add_paragraph(
        "This Agreement shall be governed by and construed in accordance with the laws of "
        "Germany, excluding its conflict of law provisions and the United Nations Convention "
        "on Contracts for the International Sale of Goods."
    )

    doc.add_heading("9. Entire Agreement", level=1)
    doc.add_paragraph(
        "This Agreement constitutes the entire agreement between the Parties concerning "
        "confidentiality with respect to the Purpose and supersedes all prior or "
        "contemporaneous oral or written understandings on the same subject. "
        "Amendments must be in writing and signed by duly authorised representatives of both Parties."
    )

    doc.add_heading("10. Signatures", level=1)
    doc.add_paragraph(
        "IN WITNESS WHEREOF, the Parties have executed this Mutual Non-Disclosure Agreement "
        "as of the date first written above."
    )
    doc.add_paragraph("\nAtos SE\nBy: ___________________________\nName:\nTitle:\nDate:")
    doc.add_paragraph("\nAutoDrive AG\nBy: ___________________________\nName:\nTitle:\nDate:")

    path = os.path.join(os.path.dirname(__file__), "nda_autodrive.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_sow_cybersecurity():
    """Reference SOW: Atos with EuroBank for a cybersecurity assessment — all required clauses present."""
    doc = Document()
    doc.add_heading("STATEMENT OF WORK", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Cybersecurity Posture Assessment and Remediation Roadmap", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Statement of Work (\"SOW\") is entered into as of 1 June 2025, between "
        "Atos IT Solutions and Services, S.A.S. (\"Service Provider\"), a company organised "
        "under the laws of France, with its registered office at 80 Quai Voltaire, 95870 "
        "Bezons, France, and EuroBank N.V. (\"Client\"), a company incorporated under the "
        "laws of the Netherlands, with its registered office at Herengracht 458, 1017 CA "
        "Amsterdam, the Netherlands. This SOW is incorporated into and governed by the Master "
        "Services Agreement between the Parties dated 15 January 2025 "
        "(MSA Reference: MSA-EUROBANK-2025-001)."
    )

    doc.add_heading("1. Project Scope", level=1)
    doc.add_paragraph(
        "Service Provider shall conduct a comprehensive cybersecurity posture assessment of "
        "Client's IT environment, covering: network infrastructure, cloud workloads "
        "(Microsoft Azure, AWS), endpoint estate, identity and access management controls, "
        "and compliance with Directive (EU) 2022/2555 (NIS2) and EBA ICT Risk Management Guidelines. "
        "The assessment shall include technical controls review, process maturity assessment "
        "(based on NIST CSF 2.0), and regulatory gap analysis."
    )
    doc.add_paragraph("The following are explicitly OUT OF SCOPE:")
    out_of_scope = [
        "Penetration testing or red team exercises (may be commissioned under a separate SOW)",
        "Security monitoring services or managed detection and response (MDR)",
        "Remediation implementation — Service Provider will advise; Client implements",
        "Assessment of third-party or supplier systems not directly integrated with Client's environment",
    ]
    for item in out_of_scope:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Deliverables and Acceptance Criteria", level=1)
    doc.add_paragraph(
        "Each deliverable is subject to written acceptance by Client within ten (10) business "
        "days of delivery. Silence for more than ten (10) business days constitutes acceptance."
    )
    deliverables = [
        ("D1", "Kick-off and Scoping Workshop", "20 June 2025",
         "Workshop held and documented; scope, access requirements, and stakeholder map confirmed in writing."),
        ("D2", "Technical Assessment Report (Draft)", "31 July 2025",
         "Draft report covering all five domains with findings, severity ratings, and evidence appendices."),
        ("D3", "Regulatory Gap Analysis", "15 August 2025",
         "Documented gaps against NIS2 and EBA ICT Guidelines, each mapped to a severity and regulatory reference."),
        ("D4", "Final Cybersecurity Posture Assessment Report", "5 September 2025",
         "Consolidated final report incorporating Client review feedback, executive summary, and full findings register."),
        ("D5", "Remediation Roadmap and Prioritisation Matrix", "26 September 2025",
         "Prioritised 12-month remediation roadmap aligned to risk severity and regulatory deadlines, accepted by CISO."),
    ]
    for code, name, due, criteria in deliverables:
        doc.add_paragraph(f"{code} — {name} (Due: {due})", style="List Bullet")
        p = doc.add_paragraph(f"Acceptance criteria: {criteria}")
        p.paragraph_format.left_indent = Pt(36)

    doc.add_heading("3. Payment Terms", level=1)
    doc.add_paragraph(
        "Total contract value: EUR 185,000 (excluding VAT). Payment schedule:"
    )
    payments = [
        "20% (EUR 37,000) — upon SOW execution",
        "30% (EUR 55,500) — upon acceptance of D2 and D3",
        "50% (EUR 92,500) — upon acceptance of D5",
    ]
    for p in payments:
        doc.add_paragraph(p, style="List Bullet")
    doc.add_paragraph(
        "All invoices are payable within thirty (30) calendar days of receipt. "
        "Late payments shall accrue interest in accordance with applicable Dutch law."
    )

    doc.add_heading("4. Roles and Responsibilities", level=1)
    doc.add_paragraph("Service Provider shall provide:")
    sp_roles = [
        "Engagement Lead: single point of accountability, responsible for quality and escalation",
        "Lead Security Architect: technical lead for assessment methodology and findings",
        "2× Security Analysts: technical data collection and testing",
        "Compliance Specialist: regulatory gap analysis (NIS2 and EBA)",
    ]
    for r in sp_roles:
        doc.add_paragraph(r, style="List Bullet")
    doc.add_paragraph("Client shall provide:")
    client_roles = [
        "Engagement Sponsor (CISO or equivalent): executive decisions and final acceptance",
        "IT Infrastructure Lead: access to technical environments and documentation",
        "Compliance Officer: input on regulatory context and existing control evidence",
        "Named availability of relevant technical staff for interview sessions (up to 20 hours total)",
    ]
    for r in client_roles:
        doc.add_paragraph(r, style="List Bullet")

    doc.add_heading("5. Change Management", level=1)
    doc.add_paragraph(
        "Any proposed change to scope, timeline, or cost must be submitted via a written "
        "Change Request (\"CR\") form. Service Provider shall assess and respond to each CR "
        "within five (5) business days with an impact assessment covering effort, schedule, and "
        "revised fees. No change to this SOW shall take effect until a CR has been approved in "
        "writing by the designated project managers of both Parties."
    )

    doc.add_heading("6. Intellectual Property", level=1)
    doc.add_paragraph(
        "All reports, roadmaps, and other deliverables created specifically for Client under "
        "this SOW (\"Custom IP\") shall be owned by Client upon receipt of full payment. "
        "Service Provider retains all rights in its pre-existing methodologies, assessment "
        "frameworks, tools, and templates (\"Background IP\"). Service Provider grants Client "
        "a non-exclusive, perpetual, royalty-free licence to use Background IP embedded in "
        "the deliverables solely for Client's internal security management purposes."
    )

    doc.add_heading("7. Data Protection", level=1)
    doc.add_paragraph(
        "To the extent that Service Provider processes personal data on behalf of Client during "
        "this engagement, Service Provider acts as Data Processor and Client acts as Data "
        "Controller. The Parties shall execute a Data Processing Agreement (DPA) as required "
        "by Regulation (EU) 2016/679 (GDPR) prior to commencement of any processing activity. "
        "Service Provider shall process personal data only as necessary for the Purpose and "
        "shall implement appropriate technical and organisational security measures. "
        "All personal data shall be deleted or returned at the end of the engagement."
    )

    doc.add_heading("8. Warranty", level=1)
    doc.add_paragraph(
        "Service Provider warrants that all deliverables shall conform to the acceptance "
        "criteria defined in Section 2 for a period of sixty (60) days following Client's "
        "written acceptance (\"Warranty Period\"). During the Warranty Period, Service "
        "Provider shall correct any material non-conformance at no additional cost to Client "
        "within ten (10) business days of written notification."
    )

    doc.add_heading("9. Limitation of Liability", level=1)
    doc.add_paragraph(
        "Each Party's total aggregate liability to the other Party under or in connection with "
        "this SOW shall not exceed the total fees paid or payable under this SOW. Neither Party "
        "shall be liable for any indirect, special, consequential, incidental, or punitive "
        "damages, including loss of profits, loss of revenue, or loss of data, even if advised "
        "of the possibility of such damages. The foregoing limitations do not apply to damages "
        "arising from fraud, wilful misconduct, or death or personal injury caused by negligence."
    )

    doc.add_heading("10. Termination", level=1)
    doc.add_paragraph(
        "Either Party may terminate this SOW for convenience by giving thirty (30) calendar "
        "days' written notice to the other Party. In the event of termination for convenience "
        "by Client, Client shall pay for all work satisfactorily completed and accepted up to "
        "the effective termination date, plus documented wind-down costs not exceeding EUR 10,000. "
        "Either Party may terminate this SOW with immediate effect by written notice if the "
        "other Party commits a material breach that remains uncured for fifteen (15) business "
        "days after written notice of the breach."
    )

    doc.add_heading("11. Force Majeure", level=1)
    doc.add_paragraph(
        "Neither Party shall be liable for any failure or delay in performing its obligations "
        "under this SOW to the extent caused by events or circumstances beyond its reasonable "
        "control, including acts of God, natural disasters, war, terrorism, industrial action, "
        "government restrictions, or widespread failures of telecommunications or internet "
        "infrastructure. The affected Party shall notify the other Party within five (5) "
        "business days of the onset of such an event and shall use reasonable efforts to "
        "mitigate its effects and resume performance as soon as practicable."
    )

    doc.add_heading("12. Governing Law and Dispute Resolution", level=1)
    doc.add_paragraph(
        "This SOW shall be governed by and construed in accordance with the laws of France. "
        "Any dispute arising out of or in connection with this SOW shall first be referred to "
        "the respective senior management of the Parties for good-faith resolution within "
        "thirty (30) days of written notice. Any unresolved dispute shall be submitted to the "
        "exclusive jurisdiction of the Commercial Court of Paris (Tribunal de Commerce de Paris)."
    )

    doc.add_heading("13. Signatures", level=1)
    doc.add_paragraph(
        "IN WITNESS WHEREOF, the Parties have executed this Statement of Work as of the date "
        "first written above."
    )
    doc.add_paragraph("\nAtos IT Solutions and Services, S.A.S.\nBy: ___________________________\nName:\nTitle:\nDate:")
    doc.add_paragraph("\nEuroBank N.V.\nBy: ___________________________\nName:\nTitle:\nDate:")

    path = os.path.join(os.path.dirname(__file__), "sow_eurobank_cybersecurity.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_hiring_data_engineer():
    """Reference Hiring contract: Senior Data Engineer at Atos SE (France) — all required clauses present."""
    doc = Document()
    doc.add_heading("EMPLOYMENT CONTRACT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Senior Data Engineer", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Employment Contract (\"Contract\") is entered into as of 1 July 2025, between:"
    )
    doc.add_paragraph(
        "Atos SE, a société européenne organised under the laws of France, with its registered "
        "office at 80 Quai Voltaire, 95870 Bezons, France (\"Employer\"), represented by "
        "Marie Leclerc, Vice President Human Resources; and"
    )
    doc.add_paragraph(
        "Léa Fontaine, born 14 March 1992, residing at 12 Rue des Lilas, 75019 Paris, France "
        "(\"Employee\")."
    )
    doc.add_paragraph(
        "Employee is engaged in the role of Senior Data Engineer (Grade: P3) within the Data "
        "and Analytics business unit, reporting to the Head of Data Engineering, EMEA. "
        "The start date is 1 August 2025. This Contract is governed by the French Labour Code "
        "(Code du travail) and the applicable collective bargaining agreement "
        "(Convention Collective Nationale des Bureaux d'Études Techniques, SYNTEC)."
    )

    doc.add_heading("1. Compensation", level=1)
    doc.add_paragraph(
        "Employer shall pay Employee a gross annual base salary of EUR 72,000, paid in "
        "twelve (12) equal monthly instalments on the last working day of each calendar month. "
        "Employee is eligible for a discretionary annual performance bonus of up to 15% of "
        "annual base salary, subject to achievement of personal and business objectives agreed "
        "annually in writing. Salary shall be reviewed no less than once per year. "
        "Employer shall comply with all applicable payroll deduction obligations including "
        "social security contributions (cotisations sociales) under French law."
    )

    doc.add_heading("2. Working Hours", level=1)
    doc.add_paragraph(
        "Employee's standard working hours are 37.5 hours per week, Monday to Friday, "
        "with a one (1) hour unpaid lunch break. Normal working hours are 09:00–17:30 at "
        "Employee's primary place of work. Employee may work up to two (2) days per week "
        "from a pre-approved remote location under Employer's Flexible Working Policy. "
        "Overtime, when required and pre-approved in writing by Employee's line manager, "
        "shall be compensated by time off in lieu in accordance with the SYNTEC agreement. "
        "Employee shall be entitled to 25 days of paid annual leave per year, plus all "
        "public holidays observed in France."
    )

    doc.add_heading("3. Probation Period", level=1)
    doc.add_paragraph(
        "Employment is subject to a probation period of three (3) months commencing on the "
        "start date, which may be extended once for a further period of three (3) months by "
        "mutual written agreement. During the probation period, either Party may terminate "
        "this Contract by giving two (2) weeks' written notice. Employer will confirm "
        "successful completion of the probation period in writing."
    )

    doc.add_heading("4. Benefits", level=1)
    doc.add_paragraph(
        "Employee shall be entitled to the following benefits from the start date, subject to "
        "the terms of the relevant scheme rules:"
    )
    benefits = [
        "Private health insurance (mutuelle d'entreprise): Employer contributes 60% of premium",
        "Group pension plan (retraite complémentaire AGIRC-ARRCO): contributions as required by law",
        "Life assurance: 4× annual base salary",
        "Meal vouchers (tickets restaurant): face value EUR 9.50 per working day, Employer contributes 60%",
        "Commuting subsidy: 50% of Navigo annual pass (Île-de-France), as required by French law",
        "Annual training budget: EUR 2,000 per calendar year for role-relevant certifications",
    ]
    for b in benefits:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5. Intellectual Property and Inventions", level=1)
    doc.add_paragraph(
        "Any and all inventions, developments, discoveries, improvements, software, data "
        "models, algorithms, or other intellectual property created or developed by Employee "
        "during the term of employment, whether alone or jointly, using Employer's resources "
        "or during working hours, shall be the exclusive property of Employer and are hereby "
        "assigned to Employer. Employee shall promptly disclose any such creation to Employer "
        "in writing and shall execute any documents reasonably requested to perfect Employer's "
        "title. Employee retains rights to intellectual property created entirely in Employee's "
        "own time, without use of Employer's resources, and unrelated to Employer's business."
    )

    doc.add_heading("6. Confidentiality", level=1)
    doc.add_paragraph(
        "Employee shall, both during the term of this Contract and indefinitely thereafter, "
        "keep strictly confidential all Confidential Information. \"Confidential Information\" "
        "means any non-public information relating to Employer, its clients, suppliers, or "
        "partners, including but not limited to: client data, pricing, technical architectures, "
        "source code, business strategies, personnel information, and financial data. "
        "Employee shall not use Confidential Information for any purpose other than the "
        "performance of duties under this Contract, and shall not disclose it to any third "
        "party without Employer's prior written consent. These obligations survive termination "
        "of employment without limit of time in respect of trade secrets, and for three (3) "
        "years for all other Confidential Information."
    )

    doc.add_heading("7. Termination and Notice Periods", level=1)
    doc.add_paragraph(
        "Following satisfactory completion of the probation period, either Party may terminate "
        "this Contract by giving three (3) months' written notice. Employer may elect to make "
        "a payment in lieu of notice (PILON). In the event of serious misconduct (faute grave), "
        "Employer may dismiss Employee with immediate effect and without payment in lieu of "
        "notice, subject to the disciplinary procedure required under French law. "
        "In the event of gross misconduct (faute lourde), no notice pay or severance shall be "
        "payable. Upon termination for any reason, Employee shall return all Employer property "
        "and delete all Employer data from personal devices."
    )

    doc.add_heading("8. Governing Law and Jurisdiction", level=1)
    doc.add_paragraph(
        "This Contract shall be governed by and construed in accordance with French law. "
        "The Parties agree that the Conseil de Prud'hommes of Paris shall have exclusive "
        "jurisdiction to resolve any dispute arising out of or in connection with this Contract."
    )

    doc.add_heading("9. Signatures", level=1)
    doc.add_paragraph(
        "This Contract is executed in two (2) original copies, one for each Party. "
        "By signing below, both Parties confirm they have read, understood, and agree to "
        "be bound by the terms of this Contract."
    )
    doc.add_paragraph("\nAtos SE\nBy: ___________________________\nName: Marie Leclerc\nTitle: VP Human Resources\nDate:")
    doc.add_paragraph("\nEmployee\nBy: ___________________________\nName: Léa Fontaine\nDate:")

    path = os.path.join(os.path.dirname(__file__), "hiring_lea_fontaine_data_engineer.docx")
    doc.save(path)
    print(f"Created: {path}")


def create_hiring_cloud_architect():
    """Reference Hiring contract: Cloud Solutions Architect at Atos SE Spain — all required clauses present."""
    doc = Document()
    doc.add_heading("EMPLOYMENT CONTRACT", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Cloud Solutions Architect", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This Employment Contract (\"Contract\") is entered into as of 15 July 2025, between:"
    )
    doc.add_paragraph(
        "Atos IT Solutions and Services, S.A.U., a company organised under the laws of Spain "
        "(NIF: A-28598071), with its registered office at Albarracín 25, 28037 Madrid, Spain "
        "(\"Employer\"), represented by Carlos Moreno, Director of Human Resources; and"
    )
    doc.add_paragraph(
        "Daniel Pérez García, with DNI 52345678-X, residing at Calle Mayor 14, 28013 Madrid, "
        "Spain (\"Employee\")."
    )
    doc.add_paragraph(
        "Employee is engaged in the role of Cloud Solutions Architect (Group: II, Level: 1) "
        "within the Cloud Infrastructure Services business unit, reporting to the Head of Cloud "
        "Architecture, Iberia. The start date is 1 September 2025. This Contract is governed by "
        "the Spanish Workers' Statute (Estatuto de los Trabajadores) and the applicable "
        "collective bargaining agreement for consulting and technology companies (Convenio "
        "Colectivo Estatal de Empresas de Consultoría, Tecnologías de la Información y Estudios "
        "de Mercado y de la Opinión Pública)."
    )

    doc.add_heading("1. Compensation", level=1)
    doc.add_paragraph(
        "Employer shall pay Employee a gross annual base salary of EUR 68,000, distributed in "
        "fourteen (14) payments: twelve (12) monthly instalments paid on the last working day "
        "of each month, plus two (2) extraordinary payments (in June and December) as required "
        "by the applicable collective bargaining agreement. Employee is eligible for a "
        "discretionary annual target incentive of up to 12% of annual base salary, subject to "
        "the achievement of personal and business objectives set and measured annually in writing."
    )

    doc.add_heading("2. Working Hours", level=1)
    doc.add_paragraph(
        "Employee's regular working hours are 40 hours per week, distributed Monday to Friday, "
        "09:00–18:00, with a one (1) hour lunch break. Employee is eligible for hybrid working: "
        "a minimum of three (3) days per week at Employer's Madrid offices, with the remaining "
        "days worked remotely, subject to Employer's Remote Work Policy. Overtime shall be "
        "authorised in advance by Employee's line manager and compensated in accordance with the "
        "applicable collective bargaining agreement and Spanish law. Employee shall be entitled "
        "to 22 business days of paid annual leave per year, in addition to Spanish public holidays."
    )

    doc.add_heading("3. Probation Period", level=1)
    doc.add_paragraph(
        "This Contract is subject to a probation period of six (6) months from the start date, "
        "in accordance with the applicable collective bargaining agreement. During the probation "
        "period, either Party may terminate this Contract without cause and without notice or "
        "indemnity obligation, except for salary and benefits accrued to the date of termination. "
        "Successful completion of the probation period will be confirmed by Employer in writing."
    )

    doc.add_heading("4. Benefits", level=1)
    doc.add_paragraph(
        "From the start date, Employee shall be entitled to the following benefits under the "
        "applicable scheme rules:"
    )
    benefits = [
        "Private health and dental insurance: fully funded by Employer for Employee; "
        "Employee may add dependants at group rates",
        "Company pension plan: Employer contributes 3% of gross annual salary to a designated "
        "pension plan; Employee may make voluntary additional contributions",
        "Life and disability insurance: 3× annual base salary",
        "Restaurant card (Ticket Restaurant): EUR 220 per month, exempt from personal income tax "
        "within statutory limits",
        "Transport allowance: EUR 1,500 per year, paid monthly",
        "Annual training and certification budget: EUR 2,500 per calendar year",
        "Mobile phone and laptop provided and maintained by Employer",
    ]
    for b in benefits:
        doc.add_paragraph(b, style="List Bullet")

    doc.add_heading("5. Intellectual Property and Inventions", level=1)
    doc.add_paragraph(
        "All intellectual and industrial property rights in any work product, invention, "
        "software, architecture design, documentation, or other creation produced by Employee "
        "within the scope of employment or using Employer's resources shall vest in Employer "
        "upon creation by operation of law or by way of assignment herein. Employee "
        "irrevocably assigns to Employer all present and future rights Employee may acquire "
        "in such work product throughout the world. Employee shall, without further "
        "compensation, execute all documents and take all actions reasonably requested by "
        "Employer to record or perfect Employer's title. This clause does not affect Employee's "
        "moral rights under Spanish law to the extent those rights are non-waivable."
    )

    doc.add_heading("6. Confidentiality", level=1)
    doc.add_paragraph(
        "Employee shall maintain the strictest confidentiality regarding all non-public "
        "information relating to Employer, its clients, suppliers, and business partners "
        "that Employee accesses in the course of employment (\"Confidential Information\"). "
        "Confidential Information includes, without limitation: client identities and data, "
        "system architectures, proprietary tools and methodologies, commercial terms, and "
        "strategic plans. Employee shall not disclose Confidential Information to any person "
        "outside Employer without prior written authorisation, and shall not use it for any "
        "purpose other than the performance of duties. The obligation of confidentiality "
        "continues for five (5) years following termination of employment for all Confidential "
        "Information, and indefinitely for trade secrets."
    )

    doc.add_heading("7. Termination and Notice Periods", level=1)
    doc.add_paragraph(
        "Following the probation period, this Contract may be terminated as follows:"
    )
    termination_terms = [
        "By Employee: thirty (30) calendar days' written notice to Employer",
        "By Employer for objective reasons (causas objetivas): in accordance with Articles "
        "51 and 52 of the Workers' Statute, with applicable notice and severance entitlements",
        "By Employer for disciplinary dismissal (despido disciplinario): immediate effect "
        "upon written notification of the grounds, in accordance with Article 54 of the Workers' Statute",
    ]
    for t in termination_terms:
        doc.add_paragraph(t, style="List Bullet")
    doc.add_paragraph(
        "In all cases, upon termination, Employee shall immediately return all Employer "
        "property and delete Employer data from personal devices. Severance entitlements, "
        "where applicable, shall be calculated in accordance with Spanish law."
    )

    doc.add_heading("8. Non-Compete", level=1)
    doc.add_paragraph(
        "Employee agrees that for a period of twelve (12) months following the termination of "
        "this Contract, Employee shall not, without Employer's prior written consent, engage "
        "in, or hold a position with, any business that directly competes with Employer's "
        "Cloud Infrastructure Services division in Spain. In consideration of this restriction, "
        "Employer shall pay Employee a monthly non-compete compensation of 30% of average "
        "monthly gross salary during the restriction period, as required by Article 21.2 of "
        "the Workers' Statute. Employer may waive this non-compete restriction at any time "
        "by written notice to Employee."
    )

    doc.add_heading("9. Governing Law and Jurisdiction", level=1)
    doc.add_paragraph(
        "This Contract is governed by and construed in accordance with the laws of Spain. "
        "Any dispute arising out of or in connection with this Contract shall be submitted "
        "to the Social Courts (Juzgados de lo Social) of Madrid, which shall have exclusive "
        "jurisdiction."
    )

    doc.add_heading("10. Signatures", level=1)
    doc.add_paragraph(
        "This Contract is executed in two (2) original copies, one for each Party. "
        "By signing below, both Parties confirm that they have read and understood this "
        "Contract and agree to be bound by its terms."
    )
    doc.add_paragraph("\nAtos IT Solutions and Services, S.A.U.\nBy: ___________________________\nName: Carlos Moreno\nTitle: Director of Human Resources\nDate:")
    doc.add_paragraph("\nEmployee\nBy: ___________________________\nName: Daniel Pérez García\nDNI: 52345678-X\nDate:")

    path = os.path.join(os.path.dirname(__file__), "hiring_daniel_perez_cloud_architect.docx")
    doc.save(path)
    print(f"Created: {path}")


if __name__ == "__main__":
    create_sample_nda()
    create_sample_sow()
    create_mclinic_sow_1()
    create_mclinic_sow_2()
    create_nda_fintech()
    create_nda_automotive()
    create_sow_cybersecurity()
    create_hiring_data_engineer()
    create_hiring_cloud_architect()
    print("Sample contracts generated successfully.")
