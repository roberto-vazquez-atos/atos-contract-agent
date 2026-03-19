"""Parse Word documents and extract raw text."""
from docx import Document
import re


def extract_text_from_docx(file_path: str) -> str:
    """Extract full text from a .docx file."""
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_text_from_uploaded_file(uploaded_file) -> str:
    """Extract full text from a Streamlit UploadedFile object."""
    doc = Document(uploaded_file)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def detect_contract_type(text: str) -> str:
    """Heuristic detection of contract type from text."""
    text_lower = text.lower()
    scores = {
        "NDA": sum([
            "non-disclosure" in text_lower,
            "confidentiality" in text_lower,
            "confidential information" in text_lower,
            "disclosing party" in text_lower,
            "receiving party" in text_lower,
        ]),
        "SOW": sum([
            "statement of work" in text_lower,
            "deliverable" in text_lower,
            "milestone" in text_lower,
            "scope of work" in text_lower,
            "acceptance criteria" in text_lower,
        ]),
        "HIRING": sum([
            "employment agreement" in text_lower,
            "employee" in text_lower,
            "salary" in text_lower,
            "compensation" in text_lower,
            "job title" in text_lower,
            "start date" in text_lower,
            "termination" in text_lower and "employment" in text_lower,
        ]),
    }
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "UNKNOWN"


def chunk_text(text: str, max_chars: int = 12000) -> list[str]:
    """Split text into chunks for large contracts."""
    if len(text) <= max_chars:
        return [text]
    chunks = []
    while text:
        chunks.append(text[:max_chars])
        text = text[max_chars:]
    return chunks
